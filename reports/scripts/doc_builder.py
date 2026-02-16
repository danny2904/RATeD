
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class DocBuilder:
    def __init__(self, filename='reports/ViXHate_Technical_Report.docx'):
        self.doc = Document()
        self.filename = filename

    def add_title_section(self):
        title = self.doc.add_heading('BÁO CÁO KỸ THUẬT: VI-XHATE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph('Dự án: Phát hiện và Giải thích Ngôn từ Thù ghét (Vi-XHate)')
        self.doc.add_paragraph('Ngày báo cáo: 15/01/2026')
        self.doc.add_paragraph('Model: PhoBERT (VN) & RoBERTa (EN) Multi-task Learning')

    def add_overview(self):
        self.doc.add_heading('1. Tổng quan Pipeline', level=1)
        self.doc.add_paragraph(
            'Hệ thống được xây dựng dựa trên kiến trúc Multi-task Learning (MTL) tổng quát. '
            '(1) Phân loại mức độ độc hại (Sequence Classification) và (2) Trích xuất từ ngữ gây thù ghét (Toxic Spans Detection).'
        )
        # Insert Architecture Image Logic (simplified path check)
        if os.path.exists('reports/figures/Vi-XHate-Architecture.png'):
            self.doc.add_picture('reports/figures/Vi-XHate-Architecture.png', width=Inches(6.0))
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_experimental_results(self):
        self.doc.add_heading('2. Kết quả Thực nghiệm (Experimental Results)', level=1)
        self.doc.add_heading('2.1. Tiếng Việt (ViHOS Dataset)', level=2)
        
        # Table Logic (Simplified hardcoded for now or pass in dicts)
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Metric'; hdr[1].text = 'Giá trị'; hdr[2].text = 'So với SOTA'
        row = table.add_row().cells; row[0].text='Accuracy'; row[1].text='0.9105'; row[2].text='+0.1105'
        row = table.add_row().cells; row[0].text='F1-Macro'; row[1].text='0.9101'; row[2].text='Robust Balanced'
        row = table.add_row().cells; row[0].text='Mean Span-IOU'; row[1].text='0.6849'; row[2].text='+0.0574 (> SOTA)'

        self.doc.add_heading('2.2. Tiếng Anh (HateXplain Dataset)', level=2)
        hx_table = self.doc.add_table(rows=1, cols=4)
        hx_table.style = 'Table Grid'
        hdr = hx_table.rows[0].cells
        hdr[0].text='Model'; hdr[1].text='Metric'; hdr[2].text='Giá trị'; hdr[3].text='So sánh SOTA'
        row = hx_table.add_row().cells; row[0].text='Vi-XHate (MTL)'; row[1].text='F1-Macro'; row[2].text='68.77%'; row[3].text='> BERT (66.0%)'


    def add_detailed_analysis_images(self, lang, file_suffix, label_list):
        # 3.1 or 3.2 logic
        # Just adds the images if they exist
        pass # Implemented in build_report main flow usually for flexibility

    def add_error_analysis(self, vh_errors, hx_errors, vh_span_errors, hx_span_errors):
        self.doc.add_heading('4. Phân tích lỗi (Error Analysis)', level=1)
        
        # VN
        self.doc.add_heading('4.1. Tiếng Việt (ViHOS)', level=2)
        self.doc.add_heading('4.1.1. Lỗi Phân loại', level=3)
        if vh_errors:
            for err in vh_errors:
                p = self.doc.add_paragraph()
                p.add_run("Text: ").bold = True; p.add_run(f"{err['text']}")
                p.add_run(f"\nTrue: {err['true']} | Pred: {err['pred']} | Conf: {err['conf']:.4f}").italic = True
        else: self.doc.add_paragraph("No detailed data.")

        self.doc.add_heading('4.1.2. Lỗi Trích xuất Span', level=3)
        if vh_span_errors:
            for err in vh_span_errors:
                p = self.doc.add_paragraph()
                p.add_run(f"[{err['type']}] ").bold=True; p.add_run(f"Text: {err['text'][:100]}...")
                p.add_run(f"\nGT: {err['gt']} | Pred: {err['pred']}").italic=True
        else: self.doc.add_paragraph("No span error data.")

        # EN
        self.doc.add_heading('4.2. Tiếng Anh (HateXplain)', level=2)
        self.doc.add_heading('4.2.1. Lỗi Phân loại', level=3)
        if hx_errors:
             for err in hx_errors:
                p = self.doc.add_paragraph()
                p.add_run("Text: ").bold = True; p.add_run(f"{err['text'][:200]}...")
                p.add_run(f"\nTrue: {err['true']} | Pred: {err['pred']} | Conf: {err['conf']:.4f}").italic = True

        self.doc.add_heading('4.2.2. Lỗi Trích xuất Span', level=3)
        if hx_span_errors:
             for err in hx_span_errors:
                p = self.doc.add_paragraph()
                p.add_run(f"[{err.get('type')}] ").bold=True; p.add_run(f"Text: {err.get('text')[:100]}...")
                p.add_run(f"\nGT: {err.get('gt')} | Pred: {err.get('pred')}").italic=True

    def save(self):
        self.doc.add_heading('5. Kết luận chung', level=1)
        self.doc.add_paragraph('Nghiên cứu khẳng định tính hiệu quả của kiến trúc Multi-task Learning.')
        
        self.doc.add_heading('6. Hướng phát triển', level=1)
        self.doc.add_paragraph('1. Integratek LLM explanation.')
        self.doc.add_paragraph('2. Multi-modal memes.')

        self.doc.add_heading('7. Tài liệu tham khảo', level=1)
        self.doc.add_paragraph('[1] HateXplain AAAI 2021.')
        self.doc.add_paragraph('[2] ViHOS EACL 2023.')
        self.doc.add_paragraph('[3] PhoBERT EMNLP 2020.')

        self.doc.save(self.filename)
        print(f"✅ Saved report to {self.filename}")
